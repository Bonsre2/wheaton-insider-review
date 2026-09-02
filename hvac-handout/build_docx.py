#!/usr/bin/env python3
"""Build the editable Word (.docx) version of the Bons Realty HVAC handout."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
A = lambda f: os.path.join(HERE, "assets", f)

NAVY   = RGBColor(0x1C, 0x3A, 0x5E)
SLATE  = RGBColor(0x5B, 0x6B, 0x7A)
INK    = RGBColor(0x26, 0x30, 0x3A)
ACCENT = RGBColor(0xB9, 0x8A, 0x3E)
NAVY_HEX, GRAY_HEX, LINE_HEX, WHITE_HEX = "1C3A5E", "F4F6F9", "D9DEE5", "FFFFFF"

FONT = "Segoe UI"

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(9)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.12

# ---- page geometry (Letter, tight margins) ----
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.45)
sec.left_margin = sec.right_margin = Inches(0.55)


# ------------------------- helpers -------------------------
def shade(el, hex_color):
    pr = el.get_or_add_pPr() if el.tag.endswith("}p") else el.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    pr.append(shd)


def cell_shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, top=40, bottom=40, left=90, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa"); m.append(e)
    tcPr.append(m)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none"); borders.append(e)
    tblPr.append(borders)


def box_borders(table, color=LINE_HEX, sz=6):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def run(p, text, size=9, bold=False, color=INK, italic=False, spacing=None, caps=False):
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        rpr = r._element.get_or_add_rPr()
        s = OxmlElement("w:spacing"); s.set(qn("w:val"), str(spacing)); rpr.append(s)
    return r


def add_hyperlink(paragraph, url, text, size=6.8, color=SLATE):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyper = OxmlElement("w:hyperlink"); hyper.set(qn("r:id"), r_id)
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT); rPr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size*2))); rPr.append(sz)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "5B6B7A"); rPr.append(col)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; r.append(t)
    hyper.append(r); paragraph._p.append(hyper)
    return hyper


def spacer(pts):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(""); r.font.size = Pt(pts)
    return p


def section_head(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)
    run(p, "—  ", size=11, bold=True, color=ACCENT)
    run(p, text, size=11.5, bold=True, color=NAVY)
    return p


# ------------------------- masthead -------------------------
mast = doc.add_table(rows=1, cols=2); no_borders(mast)
mast.columns[0].width = Inches(4.6); mast.columns[1].width = Inches(2.8)
lc, rc = mast.rows[0].cells
set_cell_margins(lc, 0, 0, 0, 0); set_cell_margins(rc, 0, 0, 0, 0)

# logo + wordmark in a nested inline layout
lp = lc.paragraphs[0]; lp.paragraph_format.space_after = Pt(0)
lr = lp.add_run(); lr.add_picture(A("logo.png"), width=Inches(0.34))
lp.add_run("  ")
run(lp, "BONS REALTY", size=15, bold=True, color=NAVY, spacing=28)
sub = lc.add_paragraph(); sub.paragraph_format.space_before = Pt(1)
run(sub, "TRUSTED LOCAL GUIDANCE", size=6.5, bold=True, color=SLATE, spacing=52)

rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; rp.paragraph_format.space_after = Pt(0)
run(rp, "LISTING PACKET · HOMEOWNER RESOURCE", size=6.6, bold=True, color=SLATE, spacing=24)
rp2 = rc.add_paragraph(); rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; rp2.paragraph_format.space_before = Pt(1)
run(rp2, "Property Improvement Brief", size=8, bold=True, color=NAVY)

# navy rule under masthead
rule = doc.add_paragraph(); rule.paragraph_format.space_before = Pt(3); rule.paragraph_format.space_after = Pt(0)
pPr = rule._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "18"); bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), NAVY_HEX)
pbdr.append(bottom); pPr.append(pbdr)

# ------------------------- title band -------------------------
spacer(4)
band = doc.add_table(rows=1, cols=1); box_borders(band, NAVY_HEX, 2)
tc = band.rows[0].cells[0]; cell_shade(tc, NAVY_HEX); set_cell_margins(tc, 150, 150, 200, 200)
e = tc.paragraphs[0]; e.paragraph_format.space_after = Pt(2)
run(e, "HEATING & COOLING · EDUCATIONAL OVERVIEW", size=6.8, bold=True, color=RGBColor(0xA9,0xC2,0xDC), spacing=26)
t = tc.add_paragraph(); t.paragraph_format.space_after = Pt(3)
run(t, "Potential HVAC Upgrade Option for This Home", size=19, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
s = tc.add_paragraph()
run(s, "Efficient whole-home heating and cooling — without existing ductwork.", size=10, color=RGBColor(0xD5,0xE2,0xEF))

# ------------------------- Considering Central Air -------------------------
section_head("Considering Central Air?")
body = doc.add_paragraph(); body.paragraph_format.space_after = Pt(2)
run(body, "Homes without existing ductwork can be difficult and costly to fit for traditional central air. A ", size=9)
run(body, "ductless mini-split system", size=9, bold=True, color=NAVY)
run(body, " may offer an alternative worth exploring. A mini-split pairs one compact outdoor unit with one or more slim indoor units, delivering both ", size=9)
run(body, "heating and cooling", size=9, bold=True, color=NAVY)
run(body, " year-round. Because refrigerant lines run through a small opening rather than bulky ducts, ", size=9)
run(body, "no existing ductwork is required", size=9, bold=True, color=NAVY)
run(body, "—which can make it a practical consideration for older homes. Whether it is the right fit here depends on the property and should be confirmed by a licensed contractor.", size=9)

# ------------------------- Benefits -------------------------
spacer(4)
benefits = [
    ("ic_noduct.png",    "No Existing Ductwork Needed", "Installs where ducts don’t exist or aren’t practical."),
    ("ic_disrupt.png",   "Minimal Disruption",          "Typically installed in a day or two with little demolition."),
    ("ic_efficient.png", "Energy Efficient",            "Inverter heat-pump technology can lower energy use."),
    ("ic_quiet.png",     "Quiet Operation",             "Indoor units run softly—often near a whisper."),
    ("ic_control.png",   "Individual Room Control",     "Set each zone’s temperature independently."),
    ("ic_ranch.png",     "Great Retrofit for Ranch Homes", "Often an excellent fit for older single-story layouts."),
]
bt = doc.add_table(rows=3, cols=2); no_borders(bt)
for i, (icon, title, desc) in enumerate(benefits):
    cell = bt.rows[i // 2].cells[i % 2]
    set_cell_margins(cell, 30, 60, 20, 120)
    inner = cell.add_table(rows=1, cols=2); no_borders(inner)
    inner.columns[0].width = Inches(0.32); inner.columns[1].width = Inches(3.1)
    icell, tcell = inner.rows[0].cells
    set_cell_margins(icell, 10, 0, 0, 40); set_cell_margins(tcell, 0, 0, 0, 0)
    ip = icell.paragraphs[0]; ip.paragraph_format.space_after = Pt(0)
    ip.add_run().add_picture(A(icon), width=Inches(0.22))
    tp = tcell.paragraphs[0]; tp.paragraph_format.space_after = Pt(0); tp.paragraph_format.line_spacing = 1.05
    run(tp, title, size=9, bold=True, color=NAVY)
    dp = tcell.add_paragraph(); dp.paragraph_format.space_after = Pt(0); dp.paragraph_format.line_spacing = 1.05
    run(dp, desc, size=8.3, color=SLATE)
    # remove default empty first paragraph of the outer cell
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)

# ------------------------- System visuals -------------------------
spacer(4)
vis = doc.add_table(rows=1, cols=3); box_borders(vis, LINE_HEX, 4)
for c, (img, cap) in zip(vis.rows[0].cells, [
        ("fig_condenser.png", "Outdoor Condenser"),
        ("fig_wallunit.png",  "Indoor Wall Unit"),
        ("fig_diagram.png",   "One Outdoor Unit · Multiple Indoor Zones")]):
    cell_shade(c, GRAY_HEX); set_cell_margins(c, 60, 50, 60, 60)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    w = Inches(1.7) if "diagram" in img else Inches(0.95)
    p.add_run().add_picture(A(img), width=w)
    cp = c.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER; cp.paragraph_format.space_after = Pt(0)
    run(cp, cap.upper(), size=6.6, bold=True, color=SLATE, spacing=12)
vis.columns[0].width = Inches(1.85); vis.columns[1].width = Inches(1.85); vis.columns[2].width = Inches(3.7)

# ------------------------- Cost table -------------------------
section_head("Typical Installed Costs — Chicago Area")
rows = [
    ("Single Zone", "One room or open area — e.g., a primary bedroom or living room", "$2,500 – $7,000"),
    ("Three Zones", "Several rooms on one system — common for a small home or main level", "$7,000 – $12,000"),
    ("Four Zones",  "Whole-home coverage across four separate areas", "$9,000 – $18,000"),
]
ct = doc.add_table(rows=4, cols=3); box_borders(ct, LINE_HEX, 6)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ct.rows[0].cells
headers = ["CONFIGURATION", "WHAT IT COMMONLY SERVES", "TYPICAL INSTALLED RANGE"]
for i, c in enumerate(hdr):
    cell_shade(c, NAVY_HEX); set_cell_margins(c, 55, 55, 110, 110)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    if i == 2: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, headers[i], size=8, bold=True, color=RGBColor(0xFF,0xFF,0xFF), spacing=8)
for ri, (z, d, price) in enumerate(rows, start=1):
    cells = ct.rows[ri].cells
    if ri % 2 == 0:
        for c in cells: cell_shade(c, GRAY_HEX)
    for c in cells: set_cell_margins(c, 60, 60, 110, 110)
    p0 = cells[0].paragraphs[0]; p0.paragraph_format.space_after = Pt(0)
    run(p0, z, size=9, bold=True, color=NAVY)
    p1 = cells[1].paragraphs[0]; p1.paragraph_format.space_after = Pt(0)
    run(p1, d, size=8.4, color=SLATE)
    p2 = cells[2].paragraphs[0]; p2.paragraph_format.space_after = Pt(0); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p2, price, size=9, bold=True, color=NAVY)
ct.columns[0].width = Inches(1.5); ct.columns[1].width = Inches(3.8); ct.columns[2].width = Inches(2.1)
note = doc.add_paragraph(); note.paragraph_format.space_before = Pt(3); note.paragraph_format.space_after = Pt(0)
run(note, "Pricing varies with equipment brand, efficiency (SEER2), line-set length, electrical work, permits, and site conditions. Figures are general estimates for illustration only, not a quote.", size=7.6, italic=True, color=SLATE)

# ------------------------- Manufacturers -------------------------
section_head("Popular Manufacturers")
mfgs = [
    ("qr_mitsubishi.png", "Mitsubishi Electric",  "mitsubishicomfort.com",        "https://www.mitsubishicomfort.com"),
    ("qr_daikin.png",     "Daikin",               "daikincomfort.com",            "https://www.daikincomfort.com"),
    ("qr_fujitsu.png",    "Fujitsu (AIRSTAGE)",   "fujitsugeneral.com/us",        "https://www.fujitsugeneral.com/us/residential/"),
    ("qr_lg.png",         "LG",                   "lg.com/us/residential-hvac",   "https://www.lg.com/us/residential-hvac"),
]
mt = doc.add_table(rows=1, cols=4); box_borders(mt, LINE_HEX, 6)
mt.alignment = WD_TABLE_ALIGNMENT.CENTER
for c, (qr, name, urltext, url) in zip(mt.rows[0].cells, mfgs):
    set_cell_margins(c, 80, 70, 60, 60)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(A(qr), width=Inches(0.72))
    np = c.add_paragraph(); np.alignment = WD_ALIGN_PARAGRAPH.CENTER; np.paragraph_format.space_after = Pt(1)
    run(np, name, size=8.6, bold=True, color=NAVY)
    up = c.add_paragraph(); up.alignment = WD_ALIGN_PARAGRAPH.CENTER; up.paragraph_format.space_after = Pt(0)
    add_hyperlink(up, url, urltext, size=6.9)
for col in mt.columns: col.width = Inches(1.85)
hint = doc.add_paragraph(); hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
hint.paragraph_format.space_before = Pt(3); hint.paragraph_format.space_after = Pt(0)
run(hint, "Scan a code (or click a link) to explore each manufacturer’s official ductless line-up. Listed alphabetically by category; not an endorsement.", size=7.4, color=SLATE)

# ------------------------- Disclaimer -------------------------
spacer(6)
dt = doc.add_table(rows=1, cols=1); box_borders(dt, LINE_HEX, 6)
dc = dt.rows[0].cells[0]; cell_shade(dc, GRAY_HEX); set_cell_margins(dc, 90, 90, 130, 130)
# navy left accent border on the cell
tcPr = dc._tc.get_or_add_tcPr(); tcB = OxmlElement("w:tcBorders")
left = OxmlElement("w:left"); left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "22"); left.set(qn("w:space"), "0"); left.set(qn("w:color"), NAVY_HEX)
tcB.append(left); tcPr.append(tcB)
dh = dc.paragraphs[0]; dh.paragraph_format.space_after = Pt(2)
run(dh, "IMPORTANT DISCLAIMER", size=7, bold=True, color=NAVY, spacing=16)
db = dc.add_paragraph(); db.paragraph_format.space_after = Pt(0); db.paragraph_format.line_spacing = 1.18
run(db, "This information is provided solely for general educational purposes to illustrate one potential HVAC upgrade option. Bons Realty and its agents are not licensed HVAC contractors and make no representation or warranty regarding suitability, installation feasibility, costs, savings, permits, or code compliance. Buyers should obtain an independent evaluation and written estimate from a qualified, licensed HVAC contractor.", size=7.2, color=SLATE)

# ------------------------- References -------------------------
spacer(4)
rh = doc.add_paragraph(); rh.paragraph_format.space_after = Pt(1)
run(rh, "SOURCES & REFERENCES", size=6.8, bold=True, color=SLATE, spacing=14)
refs = [
    "1. U.S. Dept. of Energy, “Ductless Mini-Split Heat Pumps,” energy.gov — system operation, efficiency, and zoning.",
    "2. EnergySage, “How Much Does a Mini-Split Cost?” (2025) — single-zone installed-cost ranges.",
    "3. HomeGuide & Angi mini-split cost guides (2025–2026) — multi-zone installed-cost ranges.",
    "4. Carrier / manufacturer cost guides — per-zone pricing and installation factors.",
    "5. Manufacturer sites: mitsubishicomfort.com, daikincomfort.com, fujitsugeneral.com, lg.com.",
    "6. ENERGY STAR, “Ductless Heating & Cooling,” energystar.gov — efficiency & comfort benefits.",
]
rt = doc.add_table(rows=3, cols=2); no_borders(rt)
for i, ref in enumerate(refs):
    c = rt.rows[i // 2].cells[i % 2]; set_cell_margins(c, 6, 6, 0, 120)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
    run(p, ref, size=6.7, color=SLATE)
    if c.paragraphs and len(c.paragraphs) > 1:
        pass

# ------------------------- Colophon -------------------------
crule = doc.add_paragraph(); crule.paragraph_format.space_before = Pt(5); crule.paragraph_format.space_after = Pt(3)
pPr = crule._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
top = OxmlElement("w:top"); top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "6"); top.set(qn("w:space"), "1"); top.set(qn("w:color"), LINE_HEX)
pbdr.append(top); pPr.append(pbdr)
cf = doc.add_table(rows=1, cols=2); no_borders(cf)
cf.columns[0].width = Inches(3.7); cf.columns[1].width = Inches(3.7)
lc2, rc2 = cf.rows[0].cells
lp2 = lc2.paragraphs[0]; lp2.paragraph_format.space_after = Pt(0)
run(lp2, "BONS REALTY", size=7.4, bold=True, color=NAVY, spacing=8)
run(lp2, "   ·   sean@bonsrealty.com", size=7.4, color=SLATE)
rp3 = rc2.paragraphs[0]; rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT; rp3.paragraph_format.space_after = Pt(0)
run(rp3, "Educational homeowner resource · Not an offer, appraisal, or contractor estimate · Rev. 2026", size=6.8, color=SLATE)

out = os.path.join(HERE, "Bons-Realty-HVAC-Mini-Split-Handout.docx")
doc.save(out)
print("DOCX written:", out)
